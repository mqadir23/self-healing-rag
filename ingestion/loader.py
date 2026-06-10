"""
loader.py — Multi-format document loader for the Self-Healing RAG pipeline.

Supports: PDF (.pdf), Word (.docx), Plain Text (.txt), Markdown (.md)
Extracts text with rich metadata (filename, page, file type, timestamp)
for downstream chunking and indexing.
"""

import os
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from datetime import datetime, timezone
from dataclasses import dataclass, field


@dataclass
class LoadedDocument:
    """Represents a block of text extracted from a document with metadata."""
    content: str
    metadata: dict = field(default_factory=dict)


SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}


def _load_pdf(file_path: str) -> list[LoadedDocument]:
    """Extract text from a PDF file, one LoadedDocument per page."""
    docs = []
    pdf = fitz.open(file_path)
    for page_num in range(len(pdf)):
        page = pdf[page_num]
        text = page.get_text("text")
        if text.strip():
            docs.append(LoadedDocument(
                content=text,
                metadata={
                    "source": os.path.basename(file_path),
                    "file_type": "pdf",
                    "page": page_num + 1,
                    "total_pages": len(pdf),
                    "loaded_at": datetime.now(timezone.utc).isoformat(),
                }
            ))
    pdf.close()
    return docs


def _load_text(file_path: str, file_type: str) -> list[LoadedDocument]:
    """Extract text from a plain text or markdown file."""
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()

    if not text.strip():
        return []

    return [LoadedDocument(
        content=text,
        metadata={
            "source": os.path.basename(file_path),
            "file_type": file_type,
            "page": 1,
            "total_pages": 1,
            "loaded_at": datetime.now(timezone.utc).isoformat(),
        }
    )]


def _load_docx(file_path: str) -> list[LoadedDocument]:
    """Extract text from a Word (.docx) file."""
    doc = DocxDocument(file_path)
    full_text = "\n".join(
        paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()
    )

    if not full_text.strip():
        return []

    return [LoadedDocument(
        content=full_text,
        metadata={
            "source": os.path.basename(file_path),
            "file_type": "docx",
            "page": 1,
            "total_pages": 1,
            "loaded_at": datetime.now(timezone.utc).isoformat(),
        }
    )]


def load_documents(data_dir: str = "data") -> list[LoadedDocument]:
    """
    Load all supported documents from the given directory.

    Scans the directory for files with supported extensions (.pdf, .txt, .md, .docx),
    extracts text content with metadata, and returns a flat list of LoadedDocument objects.

    Args:
        data_dir: Path to the directory containing documents. Defaults to "data".

    Returns:
        A list of LoadedDocument objects with content and metadata.
    """
    all_docs = []

    if not os.path.isdir(data_dir):
        print(f"[Loader] Warning: Data directory '{data_dir}' not found.")
        return all_docs

    files = sorted(os.listdir(data_dir))
    loaded_count = 0

    for filename in files:
        file_path = os.path.join(data_dir, filename)
        ext = os.path.splitext(filename)[1].lower()

        if ext not in SUPPORTED_EXTENSIONS:
            continue

        try:
            if ext == ".pdf":
                docs = _load_pdf(file_path)
            elif ext == ".docx":
                docs = _load_docx(file_path)
            elif ext in (".txt", ".md"):
                file_type = "txt" if ext == ".txt" else "markdown"
                docs = _load_text(file_path, file_type)
            else:
                continue

            all_docs.extend(docs)
            loaded_count += 1
            print(f"[Loader] Loaded '{filename}' — {len(docs)} block(s)")

        except Exception as e:
            print(f"[Loader] Error loading '{filename}': {e}")

    print(f"[Loader] Finished: {loaded_count} file(s), {len(all_docs)} total block(s)")
    return all_docs
